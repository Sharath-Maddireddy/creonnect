import sys
from docx import Document
from docx.shared import Pt
import re

def create_docx(md_path, docx_path):
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    code_content = []
    
    for line in lines:
        line = line.rstrip() # keep spaces, strip newlines
        
        if line.startswith('```'):
            if in_code_block:
                in_code_block = False
                p = doc.add_paragraph('\n'.join(code_content))
                p.style = 'Quote' # simple enough for code
                code_content = []
            else:
                in_code_block = True
            continue
            
        if in_code_block:
            code_content.append(line)
            continue
            
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], 0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 2)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. ') or line.startswith('4. ') or line.startswith('5. '):
            doc.add_paragraph(line[3:], style='List Number')
        elif line.strip() == '' or line.startswith('---'):
            continue # ignore horizontal rules or empty lines
        elif line.startswith('|'):
            # naive table processing, just write lines
            doc.add_paragraph(line)
        else:
            # Bold parsing (naive)
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    doc.save(docx_path)

if __name__ == "__main__":
    if len(sys.argv) > 2:
        create_docx(sys.argv[1], sys.argv[2])
    else:
        create_docx("docs/Creator_discovery_architecture.md", "docs/Creator_discovery_architecture.docx")
